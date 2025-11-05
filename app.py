from flask import Flask, request, jsonify
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
from docx import Document
from datetime import datetime
import os

app = Flask(__name__)

# ---- ZOHO MAIL CONFIG ----
ZOHO_EMAIL = "arnavbhandari2328@zohomail.in"  # Replace with your Zoho Mail ID
ZOHO_APP_PASSWORD = "apcBqixT3DZW"  # Paste the app password you generated
SMTP_SERVER = "smtp.zoho.in"
SMTP_PORT = 465  # SSL port for Zoho Mail

# ---- WHATSAPP CONFIG ----
WHATSAPP_PHONE = "919009003996"  # Replace with your WhatsApp number if needed


@app.route('/')
def home():
    return jsonify({"service": "quotation-bot", "status": "ok"})


@app.route('/webhook', methods=['POST'])
def webhook():
    try:
        data = request.json
        print("Webhook received:", data)

        # Extract data from the message
        text = data.get("message", "")
        if not text:
            return jsonify({"status": "error", "message": "No text received"}), 400

        # Simple parsing logic (example: "quote 101 for Vedant, 5 pipes at 2500 per pipe")
        quote_id = "110"
        customer_name = "Vedant"
        product = "SS 316L Pipe"
        quantity = 5
        rate = 2500
        email = "vip.vedant3@gmail.com"

        # Generate Quotation Document
        file_path = generate_quotation(quote_id, customer_name, product, quantity, rate)
        print(f"Quotation file created: {file_path}")

        # Send Email via Zoho
        send_email(email, file_path, customer_name)

        print(f"Quotation successfully sent to {email}")
        return jsonify({"status": "success", "message": "Quotation sent!"})

    except Exception as e:
        print("Error in webhook:", str(e))
        return jsonify({"status": "error", "message": str(e)}), 500


def generate_quotation(quote_id, customer_name, product, quantity, rate):
    document = Document()
    document.add_heading('Quotation', level=1)

    document.add_paragraph(f"Quotation No: {quote_id}")
    document.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")
    document.add_paragraph(f"Customer Name: {customer_name}")
    document.add_paragraph(f"Product: {product}")
    document.add_paragraph(f"Quantity: {quantity}")
    document.add_paragraph(f"Rate: ₹{rate}")
    document.add_paragraph(f"Total: ₹{rate * quantity}")
    document.add_paragraph("\nThank you for your business!\nNIVEE METAL PRODUCTS PVT LTD")

    file_path = f"/tmp/Quotation_{customer_name}_{datetime.now().strftime('%Y-%m-%d')}.docx"
    document.save(file_path)
    return file_path


def send_email(recipient, attachment_path, customer_name):
    try:
        msg = MIMEMultipart()
        msg['From'] = ZOHO_EMAIL
        msg['To'] = recipient
        msg['Subject'] = f"Quotation from NIVEE METAL PRODUCTS PVT LTD (Ref: {customer_name})"

        # Email body
        body = MIMEText(
            f"Dear {customer_name},\n\nPlease find attached the quotation as requested.\n\nRegards,\nNIVEE METAL PRODUCTS PVT LTD",
            'plain'
        )
        msg.attach(body)

        # Attachment
        with open(attachment_path, "rb") as f:
            part = MIMEApplication(f.read(), Name=os.path.basename(attachment_path))
            part['Content-Disposition'] = f'attachment; filename="{os.path.basename(attachment_path)}"'
            msg.attach(part)

        # Send Email
        with smtplib.SMTP_SSL(SMTP_SERVER, SMTP_PORT) as server:
            server.login(ZOHO_EMAIL, ZOHO_APP_PASSWORD)
            server.send_message(msg)

        print(f"✅ Email sent successfully to {recipient}")

    except Exception as e:
        print(f"❌ Failed to send email: {e}")


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
